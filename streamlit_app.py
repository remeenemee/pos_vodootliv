import streamlit as st
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# "Программа предназначена для расчета притока воды в котлован"
st.set_page_config(layout="wide")
st.title("💧 Расчет притока воды в котлован")
st.markdown("На основе методики из расчетного примера и СП 103.13330.2012")

# --- Вкладки ---
tab1, tab2, tab3 = st.tabs(["Расчет", "Методика", "Экспорт в Word"])

# --- Хранение данных ---
calc_data = {}

with tab1:
    st.header("🔧 Ввод данных (отметки от поверхности земли = 0.00 м)")

    col1, col2 = st.columns(2)

    with col1:
        L = st.number_input("Длина котлована L, м", min_value=1.0, value=21.48)
        B = st.number_input("Ширина котлована B, м", min_value=1.0, value=29.2)
        depth = st.number_input("Глубина котлована, м", min_value=0.1, value=2.25)
        z_bottom = -depth  # дно котлована

    with col2:
        z_water_depth = st.number_input("Глубина УГВ от поверхности, м", min_value=0.0, value=1.10)
        z_water = -z_water_depth  # отметка УГВ
        k = st.number_input("Коэффициент фильтрации k, м/сут", min_value=0.01, value=2.0)
        reserve = st.number_input("Запас понижения ниже дна, м", min_value=0.0, value=1.0)

    # --- Выбор типа котлована ---
    st.subheader("Тип котлована")
    type_option = st.radio(
        "Выберите тип",
        ["Несовершенный", "Совершенный"],
        horizontal=True
    )

    # --- Общий расчет понижения s ---
    s = depth + reserve - z_water_depth  # s = глубина + запас - глубина УГВ
    st.info(f"🔹 Глубина понижения грунтовых вод: **s = {s:.2f} м**")

    # --- Расчет для НЕСОВЕРШЕННОГО котлована (по расчет.pdf) ---
    if type_option == "Несовершенный":
        st.info("🔹 Котлован несовершенный — дно не доходит до водоупорного слоя")

        # 1. Высота активного слоя H0
        H0 = (4 / 3) * s
        st.write(f"**2. Высота активного слоя грунта, H₀:** {H0:.2f} м")

        # 2. Радиус влияния R
        R = 1.95 * s * np.sqrt(k * H0)
        st.write(f"**4. Расчётный радиус влияния, R:** {R:.2f} м")

        # 3. Коэффициент η по B/L
        ratio = B / L
        eta = 1.18  # при B/L >= 0.6
        if ratio < 0.6:
            eta_values = {0.0: 1.00, 0.2: 1.12, 0.4: 1.16, 0.6: 1.18}
            eta = np.interp(ratio, list(eta_values.keys()), list(eta_values.values()))

        # 4. Приведённый радиус r0
        r0 = 0.25 * eta * (L + B)
        st.write(f"**5. Приведённый радиус котлована, r₀:** {r0:.2f} м (η = {eta})")

        # 5. h0
        h0 = H0 - s
        st.write(f"**6. Остаточный уровень в активной зоне, h₀:** {h0:.2f} м")

        # 6. Приток Q
        numerator = 1.36 * k * (H0**2 - h0**2)
        denominator = np.log10(R + r0) - np.log10(r0)
        Q = numerator / denominator
        Q_m3h = Q / 24

        st.write(f"**7. Приток воды в котлован, Q:** **{Q:.2f} м³/сут** ({Q_m3h:.2f} м³/ч)")

        calc_data.update({
            "type": "Несовершенный",
            "steps": [
                f"Размеры: {L:.2f} м × {B:.2f} м, глубина: {depth:.2f} м",
                f"Грунт: Песок, УГВ на {z_water_depth:.2f} м ниже поверхности",
                f"Котлован несовершенный — дно не доходит до водоупора",
                f"Коэффициент фильтрации k = {k:.2f} м/сут",
                f"Новый уровень УГВ — на {reserve:.2f} м ниже дна котлована",
                f"Глубина понижения s = {depth:.2f} + {reserve:.2f} – {z_water_depth:.2f} = {s:.2f} м",
                f"Высота активного слоя H₀ = 4/3 × {s:.2f} = {H0:.2f} м",
                f"Радиус влияния R = 1.95 × {s:.2f} × √({k:.2f} × {H0:.2f}) = {R:.2f} м",
                f"Приведённый радиус r₀ = 0.25 × {eta:.2f} × ({L:.2f} + {B:.2f}) = {r0:.2f} м",
                f"Остаточный уровень h₀ = {H0:.2f} – {s:.2f} = {h0:.2f} м",
                f"Приток Q = 1.36 × {k:.2f} × ({H0:.2f}² – {h0:.2f}²) / (lg({R:.2f} + {r0:.2f}) – lg({r0:.2f})) = {Q:.2f} м³/сут"
            ],
            "Q": Q,
            "Q_m3h": Q_m3h,
            "reserve_flow": Q_m3h * 1.3
        })

    # --- Расчет для СОВЕРШЕННОГО котлована (по СП 103, схема 8) ---
    else:
        st.info("🔹 Котлован совершенный — дно доходит до водоупорного слоя")
        z_aquiclude = st.number_input("Отметка водоупора, м", value=-5.0)

        H = abs(z_water) - z_aquiclude  # мощность водоносного пласта
        h = H - s  # уровень после понижения

        # Радиус влияния по Зихардта (с ограничением)
        R_raw = 3000 * s * np.sqrt(k)
        R = min(R_raw, 500)
        st.write(f"**Радиус влияния R:** {R:.2f} м (расчётный: {R_raw:.2f} м)")

        # Приведённый радиус
        A = L * B
        r0 = np.sqrt(A / np.pi)
        st.write(f"**Приведённый радиус r₀:** {r0:.2f} м")

        # Приток (Схема 8, СП 103)
        Q = (np.pi * k * (H**2 - h**2)) / np.log(R / r0)
        Q_m3h = Q / 24

        st.write(f"**Мощность водоносного пласта H:** {H:.2f} м")
        st.write(f"**Приток воды Q:** **{Q:.2f} м³/сут** ({Q_m3h:.2f} м³/ч)")

        calc_data.update({
            "type": "Совершенный",
            "steps": [
                f"Размеры: {L:.2f} м × {B:.2f} м, глубина: {depth:.2f} м",
                f"Грунт: Песок, УГВ на {z_water_depth:.2f} м ниже поверхности",
                f"Котлован совершенный — дно доходит до водоупора ({z_aquiclude:.2f} м)",
                f"Коэффициент фильтрации k = {k:.2f} м/сут",
                f"Глубина понижения s = {s:.2f} м",
                f"Мощность водоносного пласта H = {abs(z_water):.2f} – ({z_aquiclude:.2f}) = {H:.2f} м",
                f"Радиус влияния R = 3000 × {s:.2f} × √{k:.2f} = {R_raw:.2f} м → принят {R:.2f} м",
                f"Приведённый радиус r₀ = √({A:.2f}/π) = {r0:.2f} м",
                f"Приток Q = π × {k:.2f} × ({H:.2f}² – {h:.2f}²) / ln({R:.2f}/{r0:.2f}) = {Q:.2f} м³/сут"
            ],
            "Q": Q,
            "Q_m3h": Q_m3h,
            "reserve_flow": Q_m3h * 1.3
        })

    # --- Подбор насоса ---
    st.write("---")
    st.subheader("🛠️ Подбор насоса")
    required_flow = calc_data["Q_m3h"]
    reserved_flow = calc_data["reserve_flow"]
    st.write(f"🔹 **Требуемая производительность:** {required_flow:.2f} м³/ч")
    st.write(f"🔹 **С запасом 1.3:** **{reserved_flow:.2f} м³/ч**")

    if reserved_flow < 6:
        st.success("✅ Подойдёт переносной дренажный насос (например, ZUMFA Small D8)")
    else:
        st.warning("⚠️ Требуется более мощный насос или несколько единиц")

with tab2:
    st.header("📖 Методика расчета")
    st.markdown("""
    ### Несовершенный котлован (по расчету.pdf)
    1. **Понижение s** = глубина котлована + запас – глубина УГВ  
    2. **H₀** = 4/3 × s  
    3. **R** = 1.95 × s × √(k × H₀)  
    4. **r₀** = 0.25 × η × (L + B), где η зависит от B/L  
    5. **h₀** = H₀ – s  
    6. **Q** = 1.36 × k × (H₀² – h₀²) / (lg(R + r₀) – lg r₀)  

    ### Совершенный котлован (СП 103.13330.2012, схема 8)
    - **Q** = πk(H² – h²) / ln(R / r₀)  
    - **R** = 3000 × s × √k (не более 500 м)  
    - **r₀** = √(A / π)  

    ### Отметки
    - Поверхность земли = **0.00 м**  
    - УГВ на 1.1 м ниже → **-1.10 м**  
    - Дно котлована на 2.25 м → **-2.25 м**  
    """)

with tab3:
    st.header("📥 Экспорт в Word — как в примере 'расчет.pdf'")

    if not calc_data:
        st.warning("Выполните расчет на вкладке 'Расчет'.")
    else:
        if st.button("📄 Сформировать отчет (по образцу)"):
            doc = Document()

            # --- Стиль ---
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

            # --- Заголовок ---
            title = doc.add_heading('Расчет притока грунтовых вод', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # --- 1. Котлован ---
            doc.add_heading('1. Котлован', level=1)
            doc.add_paragraph(f"Размеры: {L:.2f} м × {B:.2f} м")
            doc.add_paragraph(f"Глубина: {depth:.2f} м")
            doc.add_paragraph(f"Грунт: Песок, грунтовые воды располагаются на {z_water_depth:.2f} метров ниже поверхности")
            doc.add_paragraph(f"Котлован {calc_data['type'].lower()}, так как его дно не доходит до водоупорного слоя")

            # --- Пошаговый расчет ---
            doc.add_heading(f"{calc_data['type'].capitalize()} котлован", level=1)
            for i, step in enumerate(calc_data["steps"], 1):
                p = doc.add_paragraph(step, style='List Number')

            # --- Итоговый приток ---
            doc.add_paragraph(f"Приток воды в котлован Q = {calc_data['Q']:.2f} м³/сут или около {calc_data['Q_m3h']:.2f} м³/ч.")

            # --- Подбор насоса ---
            doc.add_paragraph(f"Пересчитываем с коэффициентом запаса: {calc_data['Q_m3h']:.2f} × 1.3 = {calc_data['reserve_flow']:.2f} м³/ч")

            # --- Рекомендация ---
            if calc_data['reserve_flow'] < 6:
                doc.add_paragraph("")
            else:
                doc.add_paragraph(f"Требуется насос с производительностью не менее {calc_data['reserve_flow']:.2f} м³/ч.")

            # --- Сохранение ---
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="⬇️ Скачать отчет (Word .docx)",
                data=buffer,
                file_name="Расчет_притока_воды_по_образцу.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.success("✅ Отчет сформирован в точном соответствии с образцом!")

# --- Запуск: streamlit run app.py